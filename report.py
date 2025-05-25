import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
import io
import pytz
from matplotlib.ticker import MaxNLocator
import uuid
from github import Github
import base64
import os
import time

# Set seaborn style for lightweight charts
sns.set_style("whitegrid")
plt.rcParams['font.size'] = 8
plt.rcParams['axes.titlesize'] = 10
plt.rcParams['axes.labelsize'] = 8
plt.rcParams['xtick.labelsize'] = 7
plt.rcParams['ytick.labelsize'] = 7

# Set page config
st.set_page_config(page_title="Insident Verslag", layout="wide")

# Simplified CSS for mobile optimization
st.markdown("""
    <style>
        /* General layout */
        .stApp {
            background-color: #f8f9fa;
            font-family: 'Arial', sans-serif;
            color: #212529;
            text-align: center;
        }
        [data-baseweb="baseweb"] {
            background-color: #f8f9fa !important;
        }

        /* Dark mode */
        [data-theme="dark"] .stApp, [data-theme="dark"] [data-baseweb="baseweb"] {
            background-color: #212529 !important;
            color: #e9ecef !important;
        }
        [data-theme="dark"] .main .block-container {
            background-color: #343a40 !important;
            color: #e9ecef !important;
        }

        /* Main content */
        .main .block-container {
            padding: 1rem;
            background-color: #ffffff;
            border-radius: 8px;
            margin: 1rem auto;
            max-width: 1000px;
        }

        /* Headers */
        h1 {
            color: #003087;
            font-size: 2rem;
            font-weight: 600;
            margin-bottom: 1rem;
        }
        h2 {
            color: #003087;
            font-size: 1.5rem;
            font-weight: 500;
            margin: 1rem 0;
        }
        h3 {
            color: #003087;
            font-size: 1.2rem;
            font-weight: 500;
        }

        /* Input labels */
        .input-label {
            color: #003087;
            font-size: 1rem;
            font-weight: 500;
            margin-bottom: 0.3rem;
        }
        [data-theme="dark"] .input-label {
            color: #e9ecef !important;
        }

        /* Buttons */
        .stButton>button, .stDownloadButton>button {
            background-color: #003087;
            color: #ffffff !important;
            border-radius: 6px;
            padding: 0.5rem 1rem;
            font-size: 0.9rem;
            width: 100%;
            max-width: 180px;
            margin: 0.3rem auto;
            display: block;
        }
        .stButton>button:hover, .stDownloadButton>button:hover {
            background-color: #00205b;
        }
        .stDownloadButton>button {
            background-color: #007bff;
        }
        .stDownloadButton>button:hover {
            background-color: #0056b3;
        }

        /* Selectbox */
        .stSelectbox {
            background-color: #ffffff;
            border: 1px solid #ced4da;
            border-radius: 6px;
            padding: 0.4rem;
            font-size: 0.9rem;
            max-width: 100%;
            margin: 0.3rem auto;
            display: block;
        }
        [data-theme="dark"] .stSelectbox {
            background-color: #495057;
            border: 1px solid #6c757d;
            color: #e9ecef;
        }

        /* Text area */
        .stTextArea {
            background-color: #ffffff;
            border: 1px solid #ced4da;
            border-radius: 6px;
            padding: 0.4rem;
            font-size: 0.9rem;
            max-width: 100%;
            margin: 0.3rem auto;
            display: block;
        }
        [data-theme="dark"] .stTextArea {
            background-color: #495057;
            border: 1px solid #6c757d;
            color: #e9ecef;
        }

        /* Dataframe */
        .stDataFrame {
            border: 1px solid #dee2e6;
            border-radius: 6px;
            background-color: #ffffff;
            max-width: 100%;
            margin: 0 auto;
        }
        .stDataFrame th {
            background-color: #e9ecef;
            color: #003087;
            padding: 0.5rem;
            font-size: 0.85rem;
        }
        .stDataFrame td {
            padding: 0.5rem;
            font-size: 0.85rem;
        }
        [data-theme="dark"] .stDataFrame th {
            background-color: #495057;
            color: #e9ecef;
        }
        [data-theme="dark"] .stDataFrame td {
            color: #e9ecef;
        }

        /* Tabs */
        .stTabs [data-baseweb="tab"] {
            background-color: #e9ecef;
            border-radius: 6px 6px 0 0;
            padding: 0.5rem 1rem;
            font-size: 0.9rem;
            color: #003087;
            margin-right: 0.3rem;
        }
        .stTabs [data-baseweb="tab"][aria-selected="true"] {
            background-color: #003087;
            color: #ffffff;
        }
        [data-theme="dark"] .stTabs [data-baseweb="tab"] {
            background-color: #343a40;
            color: #e9ecef;
        }
        [data-theme="dark"] .stTabs [data-baseweb="tab"][aria-selected="true"] {
            background-color: #007bff;
        }

        /* Charts */
        .stPyplot {
            border-radius: 6px;
            padding: 0.5rem;
            background-color: #ffffff;
            max-width: 100%;
            margin: 0 auto;
        }
        [data-theme="dark"] .stPyplot {
            background-color: #343a40;
        }

        /* Separator */
        .custom-divider {
            border-top: 2px solid #dee2e6;
            margin: 1.5rem auto;
            max-width: 90%;
        }
        [data-theme="dark"] .custom-divider {
            border-top: 2px solid #6c757d;
        }

        /* Pagination form */
        .pagination-form .stForm {
            display: flex;
            justify-content: center;
            gap: 0.3rem;
        }

        /* Notification container */
        .notification-container {
            position: fixed;
            top: 0.5rem;
            right: 0.5rem;
            width: 90%;
            max-width: 250px;
            z-index: 1000;
        }

        /* Mobile optimization */
        @media (max-width: 768px) {
            .main .block-container {
                padding: 0.8rem;
                margin: 0.8rem;
            }
            h1 { font-size: 1.8rem; }
            h2 { font-size: 1.3rem; }
            h3 { font-size: 1.1rem; }
            .input-label { font-size: 0.9rem; }
            .stButton>button, .stDownloadButton>button {
                padding: 0.4rem 0.8rem;
                font-size: 0.85rem;
                max-width: 160px;
            }
            .stSelectbox, .stTextArea {
                font-size: 0.85rem;
                padding: 0.3rem;
                max-width: 100%;
            }
            .stDataFrame th, .stDataFrame td {
                padding: 0.4rem;
                font-size: 0.8rem;
            }
            .stTabs [data-baseweb="tab"] {
                padding: 0.4rem 0.8rem;
                font-size: 0.85rem;
            }
            .pagination-form .stForm {
                flex-direction: column;
                align-items: center;
            }
        }

        @media (max-width: 480px) {
            .main .block-container {
                padding: 0.6rem;
                margin: 0.6rem;
            }
            h1 { font-size: 1.6rem; }
            h2 { font-size: 1.2rem; }
            h3 { font-size: 1rem; }
            .stButton>button, .stDownloadButton>button {
                padding: 0.3rem 0.6rem;
                font-size: 0.8rem;
                max-width: 140px;
            }
            .stSelectbox, .stTextArea {
                font-size: 0.8rem;
                padding: 0.25rem;
            }
        }
    </style>
""", unsafe_allow_html=True)

# Mapping of incidents to categories based on the Code of Conduct
INCIDENT_TO_CATEGORY = {
    # Category 1: Minor offenses
    "Strooi van vullis": "1",
    "Eet in klas": "1",
    "Onnet voorkoms": "1",
    "Ontwrigtende gedrag in die klas": "1",
    "Wangedrag tydens samekoms": "1",
    "Betreding van verbode area": "1",
    # Category 2: Moderate offenses
    "Dros": "2",
    "Laatkom": "2",
    "Plagiaat": "2",
    "Baklei": "2",
    "Rook": "2",
    "Beskadiging van eiendom": "2",
    "Bedreiging": "2",
    # Category 3: Serious offenses
    "Boelie": "3",
    "Seksuele teistering": "3",
    "Rassistiese opmerkings": "3",
    "Dwelmverbruik": "3",
    "Onder invloed van alkohol": "3",
    "Afkyk in eksamen": "3",
    # Category 4: Severe offenses
    "Aanranding": "4",
    "Diefstal": "4",
    "Dwelmverkoop": "4",
    "Gevaarlike wapens": "4",
    "Pornografie": "4",
    "Vervalsing": "4",
    # Default for unmapped incidents (can be adjusted manually)
    "Onbekend": "1"
}

# Load and preprocess learner data
@st.cache_data
def load_learner_data():
    df = pd.read_csv("learner_list.csv")
    df.columns = df.columns.str.strip()
    df['Learner_Full_Name'] = df['Leerder van'].fillna('') + ' ' + df['Leerner se naam'].fillna('')
    df['Learner_Full_Name'] = df['Learner_Full_Name'].str.strip()
    df = df.rename(columns={
        'klasgroep': 'Class',
        'Opvoeder betrokke': 'Teacher',
        'Wat het gebeur': 'Incident',
        'Kategorie': 'Category',
        'Kommentaar': 'Comment'
    })
    df['Learner_Full_Name'] = df['Learner_Full_Name'].replace('', 'Onbekend')
    df['Class'] = df['Class'].fillna('Onbekend')
    df['Teacher'] = df['Teacher'].fillna('Onbekend')
    df['Incident'] = df['Incident'].fillna('Onbekend')
    df['Category'] = pd.to_numeric(df['Category'], errors='coerce').fillna(1).astype(int).astype(str)
    df['Comment'] = df['Comment'].fillna('Geen Kommentaar')
    np.random.seed(42)
    start_date = datetime(2024, 1, 1)
    date_range = [start_date + timedelta(days=int(x)) for x in np.random.randint(0, 365, size=len(df))]
    df['Date'] = pd.to_datetime(date_range).date
    return df

# Load or initialize incident log with Sanction_Resolved column
@st.cache_data
def load_incident_log():
    try:
        if os.path.exists("incident_log.csv") and os.path.getsize("incident_log.csv") > 0:
            df = pd.read_csv("incident_log.csv")
            if 'Learner_Name' in df.columns and 'Learner_Full_Name' not in df.columns:
                df = df.rename(columns={'Learner_Name': 'Learner_Full_Name'})
            df['Category'] = pd.to_numeric(df['Category'], errors='coerce').fillna(1).astype(int).astype(str)
            df['Class'] = df['Class'].fillna('Onbekend').astype(str)
            df['Date'] = pd.to_datetime(df['Date'], errors='coerce').dt.date
            if 'Sanction_Resolved' not in df.columns:
                df['Sanction_Resolved'] = False
            df['Sanction_Resolved'] = df['Sanction_Resolved'].astype(bool)
            return df
        else:
            return pd.DataFrame(columns=['Learner_Full_Name', 'Class', 'Teacher', 'Incident', 'Category', 'Comment', 'Date', 'Sanction_Resolved'])
    except (FileNotFoundError, pd.errors.EmptyDataError):
        return pd.DataFrame(columns=['Learner_Full_Name', 'Class', 'Teacher', 'Incident', 'Category', 'Comment', 'Date', 'Sanction_Resolved'])

# Save incident to log and push to GitHub
def save_incident(learner_full_name, class_, teacher, incident, category, comment):
    incident_log = load_incident_log()
    sa_tz = pytz.timezone('Africa/Johannesburg')
    try:
        category = str(int(float(category)))
    except ValueError:
        category = '1'
    new_incident = pd.DataFrame({
        'Learner_Full_Name': [learner_full_name],
        'Class': [class_],
        'Teacher': [teacher],
        'Incident': [incident],
        'Category': [category],
        'Comment': [comment],
        'Date': [datetime.now(sa_tz).date()],
        'Sanction_Resolved': [False]
    })
    updated_log = pd.concat([incident_log, new_incident], ignore_index=True)
    updated_log.to_csv("incident_log.csv", index=False)

    try:
        g = Github(st.secrets["GITHUB_TOKEN"])
        repo = g.get_repo("arnoldtRealph/insident")
        with open("incident_log.csv", "rb") as file:
            content = file.read()
        repo_path = "incident_log.csv"
        try:
            contents = repo.get_contents(repo_path, ref="master")
            repo.update_file(
                path=repo_path,
                message="Updated incident_log.csv with new incident",
                content=content,
                sha=contents.sha,
                branch="master"
            )
        except:
            repo.create_file(
                path=repo_path,
                message="Created incident_log.csv with new incident",
                content=content,
                branch="master"
            )
    except Exception as e:
        with open("error_log.txt", "a") as f:
            f.write(f"GitHub push failed: {str(e)}\n")

    return updated_log

# Mark sanction as resolved and update GitHub
def resolve_sanction(learner, category):
    incident_log = load_incident_log()
    mask = (incident_log['Learner_Full_Name'] == learner) & (incident_log['Category'] == category)
    if mask.any():
        incident_log.loc[mask, 'Sanction_Resolved'] = True
        incident_log.to_csv("incident_log.csv", index=False)

        try:
            g = Github(st.secrets["GITHUB_TOKEN"])
            repo = g.get_repo("arnoldtRealph/insident")
            with open("incident_log.csv", "rb") as file:
                content = file.read()
            repo_path = "incident_log.csv"
            try:
                contents = repo.get_contents(repo_path, ref="master")
                repo.update_file(
                    path=repo_path,
                    message="Updated incident_log.csv with resolved sanction",
                    content=content,
                    sha=contents.sha,
                    branch="master"
                )
            except:
                repo.create_file(
                    path=repo_path,
                    message="Created incident_log.csv with resolved sanction",
                    content=content,
                    branch="master"
                )
        except Exception as e:
            with open("error_log.txt", "a") as f:
                f.write(f"GitHub push failed: {str(e)}\n")

    return incident_log

# Clear a single incident and push to GitHub
def clear_incident(index):
    incident_log = load_incident_log()
    if 0 <= index < len(incident_log):
        updated_log = incident_log.drop(index)
        updated_log.to_csv("incident_log.csv", index=False)

        try:
            g = Github(st.secrets["GITHUB_TOKEN"])
            repo = g.get_repo("arnoldtRealph/insident")
            with open("incident_log.csv", "rb") as file:
                content = file.read()
            repo_path = "incident_log.csv"
            try:
                contents = repo.get_contents(repo_path, ref="master")
                repo.update_file(
                    path=repo_path,
                    message="Updated incident_log.csv after clearing incident",
                    content=content,
                    sha=contents.sha,
                    branch="master"
                )
            except:
                repo.create_file(
                    path=repo_path,
                    message="Created incident_log.csv after clearing incident",
                    content=content,
                    branch="master"
                )
        except Exception as e:
            with open("error_log.txt", "a") as f:
                f.write(f"GitHub push failed: {str(e)}\n")

        return updated_log
    return incident_log

# Generate Word document
def generate_word_report(df):
    doc = Document()
    doc.add_heading('Insident Verslag', 0)

    doc.add_heading('Insident Besonderhede', level=1)
    columns_to_include = ['Learner_Full_Name', 'Class', 'Teacher', 'Incident', 'Category', 'Comment', 'Date']
    filtered_df = df[columns_to_include]
    table = doc.add_table(rows=1, cols=len(columns_to_include))
    table.style = 'Table Grid'
    for i, col in enumerate(columns_to_include):
        table.cell(0, i).text = {
            'Learner_Full_Name': 'Leerder Naam',
            'Class': 'Klas',
            'Teacher': 'Onderwyser',
            'Incident': 'Insident',
            'Category': 'Kategorie',
            'Comment': 'Kommentaar',
            'Date': 'Datum'
        }.get(col, col)
    for _, row in filtered_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns_to_include):
            if col == 'Date':
                cells[i].text = row[col].strftime("%Y-%m-%d")
            else:
                cells[i].text = str(row[col])

    doc.add_heading('Insident Analise', level=1)
    
    fig, ax = plt.subplots(figsize=(3, 2))
    category_counts = df['Category'].value_counts().sort_index()
    sns.barplot(x=category_counts.index, y=category_counts.values, ax=ax, palette='Blues')
    ax.set_title('Insidente volgens Kategorie', fontsize=10)
    ax.set_xlabel('Kategorie', fontsize=8)
    ax.set_ylabel('Aantal', fontsize=8)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.tick_params(axis='both', labelsize=7)
    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=80, bbox_inches='tight')
    plt.close()
    doc.add_picture(img_stream, width=Inches(3))

    fig, ax = plt.subplots(figsize=(3, 2))
    incident_counts = df['Incident'].value_counts().head(5)
    sns.barplot(x=incident_counts.index, y=incident_counts.values, ax=ax, palette='Blues')
    ax.set_title('Insidente volgens Tipe', fontsize=10)
    ax.set_xlabel('Insident', fontsize=8)
    ax.set_ylabel('Aantal', fontsize=8)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.tick_params(axis='x', rotation=30, labelsize=7)
    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=80, bbox_inches='tight')
    plt.close()
    doc.add_picture(img_stream, width=Inches(3))

    fig, ax = plt.subplots(figsize=(3, 2))
    teacher_counts = df['Teacher'].value_counts().head(5)
    sns.barplot(x=teacher_counts.index, y=teacher_counts.values, ax=ax, palette='Blues')
    ax.set_title('Insidente volgens Onderwyser', fontsize=10)
    ax.set_xlabel('Onderwyser', fontsize=8)
    ax.set_ylabel('Aantal', fontsize=8)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.tick_params(axis='x', rotation=30, labelsize=7)
    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=80, bbox_inches='tight')
    plt.close()
    doc.add_picture(img_stream, width=Inches(3))

    fig, ax = plt.subplots(figsize=(3, 2))
    class_counts = df['Class'].value_counts().head(5)
    sns.barplot(x=class_counts.index, y=class_counts.values, ax=ax, palette='Blues')
    ax.set_title('Insidente volgens Klas', fontsize=10)
    ax.set_xlabel('Klas', fontsize=8)
    ax.set_ylabel('Aantal', fontsize=8)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.tick_params(axis='x', rotation=30, labelsize=7)
    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=80, bbox_inches='tight')
    plt.close()
    doc.add_picture(img_stream, width=Inches(3))

    fig, ax = plt.subplots(figsize=(3, 2))
    category_counts.plot(kind='pie', ax=ax, autopct='%1.1f%%', colors=sns.color_palette('Blues'), textprops={'fontsize': 7})
    ax.set_title('Insident Verspreiding', fontsize=10)
    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=80, bbox_inches='tight')
    plt.close()
    doc.add_picture(img_stream, width=Inches(3))

    doc.add_heading('Leerders met Herhalende Insidente', level=1)
    incident_counts = df['Learner_Full_Name'].value_counts()
    high_risk_learners = incident_counts[incident_counts > 2].index
    high_risk_df = df[df['Learner_Full_Name'].isin(high_risk_learners)][columns_to_include]

    if not high_risk_df.empty:
        table = doc.add_table(rows=1, cols=len(columns_to_include))
        table.style = 'Table Grid'
        headers = ['Leerder Naam', 'Klas', 'Onderwyser', 'Insident', 'Kategorie', 'Kommentaar', 'Datum']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        for _, row in high_risk_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(columns_to_include):
                if col == 'Date':
                    cells[i].text = row[col].strftime("%Y-%m-%d")
                else:
                    cells[i].text = str(row[col])
    else:
        doc.add_paragraph("Geen leerders met herhalende insidente nie.")

    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# Generate learner-specific Word report
def generate_learner_report(df, learner_full_name, period, start_date, end_date):
    doc = Document()
    doc.add_heading(f'Insident Verslag vir {learner_full_name}', 0)
    doc.add_paragraph(f'Tydperk: {period}')
    doc.add_paragraph(f'Datum Reeks: {start_date.strftime("%Y-%m-%d")} tot {end_date.strftime("%Y-%m-%d")}')

    doc.add_heading('Insident Besonderhede', level=1)
    columns_to_include = ['Learner_Full_Name', 'Class', 'Teacher', 'Incident', 'Category', 'Comment', 'Date']
    filtered_df = df[columns_to_include]
    table = doc.add_table(rows=1, cols=len(columns_to_include))
    table.style = 'Table Grid'
    for i, col in enumerate(columns_to_include):
        table.cell(0, i).text = {
            'Learner_Full_Name': 'Leerder Naam',
            'Class': 'Klas',
            'Teacher': 'Onderwyser',
            'Incident': 'Insident',
            'Category': 'Kategorie',
            'Comment': 'Kommentaar',
            'Date': 'Datum'
        }.get(col, col)
    for _, row in filtered_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns_to_include):
            if col == 'Date':
                cells[i].text = row[col].strftime("%Y-%m-%d")
            else:
                cells[i].text = str(row[col])

    if not df.empty:
        doc.add_heading('Insident Analise', level=1)
        fig, ax = plt.subplots(figsize=(3, 2))
        category_counts = df['Category'].value_counts().sort_index()
        sns.barplot(x=category_counts.index, y=category_counts.values, ax=ax, palette='Blues')
        ax.set_title('Insidente volgens Kategorie', fontsize=10)
        ax.set_xlabel('Kategorie', fontsize=8)
        ax.set_ylabel('Aantal', fontsize=8)
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        ax.tick_params(axis='both', labelsize=7)
        plt.tight_layout()
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=80, bbox_inches='tight')
        plt.close()
        doc.add_picture(img_stream, width=Inches(3))

    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# Load data
learner_df = load_learner_data()
incident_log = load_incident_log()

# Main content
with st.container():
    st.title("HOËRSKOOL SAUL DAMON")
    st.subheader("INSIDENT VERSLAG")

# Compute sanctions
if not incident_log.empty:
    tally_df = incident_log.pivot_table(
        index='Learner_Full_Name',
        columns='Category',
        values='Incident',
        aggfunc='count',
        fill_value=0
    )
    for cat in ['1', '2', '3', '4']:
        if cat not in tally_df.columns:
            tally_df[cat] = 0
    tally_df = tally_df[['1', '2', '3', '4']].reset_index()

    sanctions = []
    for _, row in tally_df.iterrows():
        learner = row['Learner_Full_Name']
        for cat in ['1', '2', '3', '4']:
            count = int(row[cat])
            if count > 0:
                if cat == '1' and count > 10:
                    sanction = 'Ouers moet afspraak maak met Mnr. Zealand.'
                elif cat == '2' and count > 5:
                    sanction = 'Ouers moet afspraak maak met Mnr. Zealand.'
                elif cat == '3' and count > 2:
                    sanction = 'Ouers moet afspraak maak met Mnr. Zealand.'
                elif cat == '4' and count >= 1:
                    sanction = 'Leerder moet geskors word.'
                else:
                    continue
                mask = (incident_log['Learner_Full_Name'] == learner) & (incident_log['Category'] == cat)
                if not incident_log[mask]['Sanction_Resolved'].all():
                    sanctions.append({
                        'Learner': learner,
                        'Category': cat,
                        'Count': count,
                        'Sanction': sanction
                    })

    sanctions_df = pd.DataFrame(sanctions)

    with st.container():
        st.markdown('<div class="notification-container">', unsafe_allow_html=True)
        any_notifications = False
        for _, row in sanctions_df.iterrows():
            learner = row['Learner']
            category = row['Category']
            any_notifications = True
            st.markdown(
                f"""
                <div style='background-color: #ffe6e6; padding: 10px; border-radius: 6px; border: 1px solid #cc0000;'>
                    <h4 style='color: #cc0000; margin: 0; font-size: 1rem;'>SANKSIEMELDING</h4>
                    <p style='color: #333; margin: 3px 0; font-size: 0.85rem;'>
                        Leerder <strong>{row['Learner']}</strong>: {row['Count']} Kategorie {row['Category']} insidente. 
                        Sanksie: {row['Sanction']}
                    </p>
                </div>
                """,
                unsafe_allow_html=True
            )
            if st.button("Opgelos", key=f"sanction_resolve_{learner}_{category}"):
                incident_log = resolve_sanction(learner, category)
                st.success("Sanksie permanent opgelos!")
                st.rerun()
        if not any_notifications:
            st.markdown(
                """
                <div style='background-color: #e6f3e6; padding: 10px; border-radius: 6px; border: 1px solid #28b463;'>
                    <p style='color: #333; margin: 0; font-size: 0.85rem;'>Geen aktiewe sanksiemeldings nie.</p>
                </div>
                """,
                unsafe_allow_html=True
            )
        st.markdown('</div>', unsafe_allow_html=True)

# Report new incident
st.header("Rapporteer Nuwe Insident")
with st.container():
    st.markdown('<div class="input-label">Leerder Naam</div>', unsafe_allow_html=True)
    learner_full_name = st.selectbox("", options=['Kies'] + sorted(learner_df['Learner_Full_Name'].unique()), key="learner_full_name")
    
    st.markdown('<div class="input-label">Klas</div>', unsafe_allow_html=True)
    class_ = st.selectbox("", options=['Kies'] + sorted(learner_df['Class'].unique()), key="class")
    
    st.markdown('<div class="input-label">Onderwyser</div>', unsafe_allow_html=True)
    teacher = st.selectbox("", options=['Kies'] + sorted(learner_df['Teacher'].unique()), key="teacher")
    
    st.markdown('<div class="input-label">Insident</div>', unsafe_allow_html=True)
    incident = st.selectbox("", options=['Kies'] + sorted(learner_df['Incident'].unique()), key="incident")
    
    if incident != 'Kies':
        default_category = INCIDENT_TO_CATEGORY.get(incident, "1")
    else:
        default_category = "Kies"
    
    st.markdown('<div class="input-label">Kategorie (Outomaties Gekies, Kan Verander Word)</div>', unsafe_allow_html=True)
    category = st.selectbox(
        "",
        options=['Kies'] + sorted(learner_df['Category'].unique(), key=lambda x: int(x)),
        index=0 if default_category == "Kies" else sorted(learner_df['Category'].unique(), key=lambda x: int(x)).index(default_category) + 1,
        key="category"
    )
    
    st.markdown('<div class="input-label">Kommentaar</div>', unsafe_allow_html=True)
    comment = st.text_area("", placeholder="Tik hier...", key="comment")
    
    if st.button("Stoor Insident"):
        if learner_full_name != 'Kies' and class_ != 'Kies' and teacher != 'Kies' and incident != 'Kies' and category != 'Kies' and comment:
            incident_log = save_incident(learner_full_name, class_, teacher, incident, category, comment)
            st.cache_data.clear()
            incident_log = load_incident_log()
            st.success("Insident suksesvol gestoor!")
            st.rerun()
        else:
            st.error("Vul asseblief alle velde in.")

# Generate learner report
st.header("Genereer Leerder Verslag")
with st.container():
    st.markdown('<div class="input-label">Kies Leerder vir Verslag</div>', unsafe_allow_html=True)
    learner_report_name = st.selectbox("", options=['Kies'] + sorted(incident_log['Learner_Full_Name'].unique()) if not incident_log.empty else ['Kies'], key="learner_report_name")
    
    st.markdown('<div class="input-label">Kies Tydperk</div>', unsafe_allow_html=True)
    report_period = st.selectbox("", options=['Daagliks', 'Weekliks', 'Maandelik', 'Kwartaalliks'], key="report_period")

    sa_tz = pytz.timezone('Africa/Johannesburg')
    today = datetime.now(sa_tz).date()

    if report_period == 'Daagliks':
        start_date = today
        end_date = today
    elif report_period == 'Weekliks':
        start_date = today - timedelta(days=today.weekday())
        end_date = start_date + timedelta(days=6)
    elif report_period == 'Maandelik':
        start_date = today.replace(day=1)
        end_date = (start_date + timedelta(days=31)).replace(day=1) - timedelta(days=1)
    else:
        quarter_start_month = ((today.month - 1) // 3) * 3 + 1
        start_date = today.replace(month=quarter_start_month, day=1)
        end_date = (start_date + timedelta(days=92)).replace(day=1) - timedelta(days=1)

    st.write(f"Verslag Datum Reeks: {start_date.strftime('%Y-%m-%d')} tot {end_date.strftime('%Y-%m-%d')}")

    if st.button("Genereer Leerder Verslag"):
        if learner_report_name != 'Kies':
            learner_incidents = incident_log[
                (incident_log['Learner_Full_Name'] == learner_report_name) &
                (incident_log['Date'] >= start_date) &
                (incident_log['Date'] <= end_date)
            ]
            if not learner_incidents.empty:
                report_stream = generate_learner_report(learner_incidents, learner_report_name, report_period, start_date, end_date)
                st.success(f"Verslag vir {learner_report_name} suksesvol gegenereer!")
                st.download_button(
                    label="Laai Leerder Verslag af",
                    data=report_stream,
                    file_name=f"insident_verslag_{learner_report_name}_{report_period.lower()}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error(f"Geen insidente gevind vir {learner_report_name} in die geselekteerde tydperk.")
        else:
            st.error("Kies asseblief 'n leerder.")

st.markdown('<div class="custom-divider"></div>', unsafe_allow_html=True)

# Incident log
st.subheader("Insident Log")
if not incident_log.empty:
    rows_per_page = 10
    total_rows = len(incident_log)
    total_pages = (total_rows + rows_per_page - 1) // rows_per_page

    if 'incident_log_page' not in st.session_state:
        st.session_state.incident_log_page = 1

    with st.form(key="pagination_form", clear_on_submit=False):
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            submit_prev = st.form_submit_button("Vorige", disabled=(st.session_state.incident_log_page <= 1))
        with col2:
            page_options = list(range(1, total_pages + 1))
            selected_page = st.selectbox("Bladsy", options=page_options, index=st.session_state.incident_log_page - 1, key="incident_log_page_select")
        with col3:
            submit_next = st.form_submit_button("Volgende", disabled=(st.session_state.incident_log_page >= total_pages))

        if submit_prev:
            st.session_state.incident_log_page = max(1, st.session_state.incident_log_page - 1)
            st.rerun()
        if submit_next:
            st.session_state.incident_log_page = min(total_pages, st.session_state.incident_log_page + 1)
            st.rerun()
        if selected_page != st.session_state.incident_log_page:
            st.session_state.incident_log_page = selected_page
            st.rerun()

    start_idx = (st.session_state.incident_log_page - 1) * rows_per_page
    end_idx = min(start_idx + rows_per_page, total_rows)

    display_df = incident_log.iloc[start_idx:end_idx].copy()
    display_df.index = range(start_idx, end_idx)

    st.dataframe(
        display_df,
        height=400,
        use_container_width=True,
        column_config={
            "Learner_Full_Name": st.column_config.TextColumn("Leerder Naam", width="medium"),
            "Class": st.column_config.TextColumn("Klas", width="small"),
            "Teacher": st.column_config.TextColumn("Onderwyser", width="medium"),
            "Incident": st.column_config.TextColumn("Insident", width="medium"),
            "Category": st.column_config.TextColumn("Kategorie", width="small"),
            "Comment": st.column_config.TextColumn("Kommentaar", width="large"),
            "Date": st.column_config.DateColumn("Datum", width="medium", format="YYYY-MM-DD"),
            "Sanction_Resolved": st.column_config.CheckboxColumn("Sanksie Opgelos", width="small")
        }
    )
    st.write(f"Wys {start_idx + 1} tot {end_idx} van {total_rows} insidente")

    st.download_button(
        label="Laai Verslag af as Word",
        data=generate_word_report(incident_log),
        file_name="insident_verslag.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.write("Verwyder 'n Insident")
    one_based_indices = list(range(1, total_rows + 1))
    st.markdown('<div class="input-label">Kies Insident om te Verwyder (deur Indeks)</div>', unsafe_allow_html=True)
    selected_display_index = st.selectbox("", options=one_based_indices, key="delete_index")
    if st.button("Verwyder Insident"):
        if 1 <= selected_display_index <= total_rows:
            zero_based_index = selected_display_index - 1
            global_index = start_idx + zero_based_index  # Adjust for paginated view
            if 0 <= global_index < len(incident_log):
                incident_log = clear_incident(global_index)
                st.cache_data.clear()
                incident_log = load_incident_log()
                st.success(f"Insident {selected_display_index} suksesvol verwyder!")
                total_rows = len(incident_log)
                total_pages = (total_rows + rows_per_page - 1) // rows_per_page
                if st.session_state.incident_log_page > total_pages and total_pages > 0:
                    st.session_state.incident_log_page = total_pages
                elif total_pages == 0:
                    st.session_state.incident_log_page = 1
                st.rerun()
            else:
                st.error("Gekose indeks is ongeldig.")
        else:
            st.error("Gekose indeks is buite bereik.")

else:
    st.write("Geen insidente in die log nie.")

# Today's incidents
st.subheader("Vandag se Insidente")
today = datetime.now(pytz.timezone('Africa/Johannesburg')).date()
today_incidents = incident_log[incident_log['Date'] == today]
if not today_incidents.empty:
    st.write(f"Totale Insidente Vandag: {len(today_incidents)}")

    st.write("Insidente volgens Kategorie")
    fig, ax = plt.subplots(figsize=(3, 2))
    category_counts = today_incidents['Category'].value_counts().sort_index()
    sns.barplot(x=category_counts.index, y=category_counts.values, ax=ax, palette='Blues')
    ax.set_title('Insidente volgens Kategorie (Vandag)', fontsize=10)
    ax.set_xlabel('Kategorie', fontsize=8)
    ax.set_ylabel('Aantal', fontsize=8)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.tick_params(axis='both', labelsize=7)
    plt.tight_layout()
    st.pyplot(fig)
    plt.close()

    st.write("Insidente volgens Tipe")
    fig, ax = plt.subplots(figsize=(6, 3))
    incident_counts = today_incidents['Incident'].value_counts().head(5)
    sns.barplot(x=incident_counts.index, y=incident_counts.values, ax=ax, palette='Blues')
    ax.set_title('Insidente volgens Tipe (Vandag)', fontsize=10)
    ax.set_xlabel('Insident', fontsize=8)
    ax.set_ylabel('Aantal', fontsize=8)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.tick_params(axis='x', rotation=30, labelsize=7)
    plt.tight_layout()
    st.pyplot(fig)
    plt.close()

    st.write("Insidente volgens Onderwyser")
    fig, ax = plt.subplots(figsize=(3, 2))
    teacher_counts = today_incidents['Teacher'].value_counts().head(5)
    sns.barplot(x=teacher_counts.index, y=teacher_counts.values, ax=ax, palette='Blues')
    ax.set_title('Insidente volgens Onderwyser (Vandag)', fontsize=10)
    ax.set_xlabel('Onderwyser', fontsize=8)
    ax.set_ylabel('Aantal', fontsize=8)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.tick_params(axis='x', rotation=30, labelsize=7)
    plt.tight_layout()
    st.pyplot(fig)
    plt.close()

    st.write("Insidente volgens Klas")
    fig, ax = plt.subplots(figsize=(3, 2))
    class_counts = today_incidents['Class'].value_counts().head(5)
    sns.barplot(x=class_counts.index, y=class_counts.values, ax=ax, palette='Blues')
    ax.set_title('Insidente volgens Klas (Vandag)', fontsize=10)
    ax.set_xlabel('Klas', fontsize=8)
    ax.set_ylabel('Aantal', fontsize=8)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.tick_params(axis='x', rotation=30, labelsize=7)
    plt.tight_layout()
    st.pyplot(fig)
    plt.close()
else:
    st.write("Geen insidente vandag gerapporteer nie.")

# Tabs for summaries
tab1, tab2, tab3, tab4 = st.tabs(["Gefiltreerde Data", "Weeklikse Opsomming", "Maandelikse Opsomming", "Kwartaallikse Opsomming"])

with tab1:
    st.subheader("Gefiltreerde Data")
    st.markdown('<div class="input-label">Filter Leerder Naam</div>', unsafe_allow_html=True)
    learner_options = ['Alle'] + sorted(incident_log['Learner_Full_Name'].unique()) if not incident_log.empty else ['Alle']
    filter_learner = st.selectbox("", options=learner_options, key="filter_learner")
    
    st.markdown('<div class="input-label">Filter Klas</div>', unsafe_allow_html=True)
    class_options = sorted(incident_log['Class'].fillna('Onbekend').astype(str).unique()) if not incident_log.empty else []
    filter_class = st.selectbox("", options=['Alle'] + class_options, key="filter_class")
    
    st.markdown('<div class="input-label">Filter Onderwyser</div>', unsafe_allow_html=True)
    teacher_options = ['Alle'] + sorted(incident_log['Teacher'].unique()) if not incident_log.empty else ['Alle']
    filter_teacher = st.selectbox("", options=teacher_options, key="filter_teacher")
    
    st.markdown('<div class="input-label">Filter Insident</div>', unsafe_allow_html=True)
    incident_options = ['Alle'] + sorted(incident_log['Incident'].unique()) if not incident_log.empty else ['Alle']
    filter_incident = st.selectbox("", options=incident_options, key="filter_incident")
    
    st.markdown('<div class="input-label">Filter Kategorie</div>', unsafe_allow_html=True)
    category_options = ['Alle'] + sorted(incident_log['Category'].unique(), key=lambda x: int(x)) if not incident_log.empty else ['Alle']
    filter_category = st.selectbox("", options=category_options, key="filter_category")
    
    filtered_df = incident_log.copy()
    if filter_learner != 'Alle':
        filtered_df = filtered_df[filtered_df['Learner_Full_Name'] == filter_learner]
    if filter_class != 'Alle':
        filtered_df = filtered_df[filtered_df['Class'] == filter_class]
    if filter_teacher != 'Alle':
        filtered_df = filtered_df[filtered_df['Teacher'] == filter_teacher]
    if filter_incident != 'Alle':
        filtered_df = filtered_df[filtered_df['Incident'] == filter_incident]
    if filter_category != 'Alle':
        filtered_df = filtered_df[filtered_df['Category'] == filter_category]
    st.dataframe(
        filtered_df.head(10),
        use_container_width=True,
        height=300,
        column_config={
            "Learner_Full_Name": st.column_config.TextColumn("Leerder Naam", width="medium"),
            "Class": st.column_config.TextColumn("Klas", width="small"),
            "Teacher": st.column_config.TextColumn("Onderwyser", width="medium"),
            "Incident": st.column_config.TextColumn("Insident", width="medium"),
            "Category": st.column_config.TextColumn("Kategorie", width="small"),
            "Comment": st.column_config.TextColumn("Kommentaar", width="large"),
            "Date": st.column_config.DateColumn("Datum", width="medium", format="YYYY-MM-DD"),
            "Sanction_Resolved": st.column_config.CheckboxColumn("Sanksie Opgelos", width="small")
        }
    )
    st.write(f"Totale Insidente: {len(filtered_df)}")

with tab2:
    st.subheader("Weeklikse Opsomming")
    if not incident_log.empty:
        incident_log_dt = incident_log.copy()
        incident_log_dt['Date'] = pd.to_datetime(incident_log_dt['Date'])
        weekly_summary = incident_log_dt.groupby([pd.Grouper(key='Date', freq='W-MON'), 'Category']).size().unstack(fill_value=0)
        weekly_summary.index = weekly_summary.index.strftime('%Y-%m-%d')
        weekly_summary['Totaal'] = weekly_summary.sum(axis=1)
        weekly_summary = weekly_summary.reset_index().rename(columns={'Date': 'Week Begin (Maandag)'})
        st.dataframe(
            weekly_summary.head(10),
            use_container_width=True,
            height=300,
            column_config={
                'Week Begin (Maandag)': st.column_config.TextColumn("Week Begin (Maandag)", width="medium"),
                'Totaal': st.column_config.NumberColumn("Totaal Insidente", width="small")
            }
        )
        fig, ax = plt.subplots(figsize=(6, 3))
        weekly_summary.set_index('Week Begin (Maandag)')[[col for col in weekly_summary.columns[1:-1]]].plot(
            kind='bar', 
            stacked=True, 
            ax=ax, 
            color=sns.color_palette('tab10', n_colors=len(weekly_summary.columns[1:-1]))
        )
        ax.set_title('Weeklikse Insidente', fontsize=10)
        ax.set_xlabel('Week Begin', fontsize=8)
        ax.set_ylabel('Aantal', fontsize=8)
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        ax.tick_params(axis='both', labelsize=7)
        ax.legend(title='Kategorie', fontsize=7)
        plt.tight_layout()
        st.pyplot(fig)
        plt.close()
    else:
        st.write("Geen insidente om te wys nie.")

with tab3:
    st.subheader("Maandelikse Opsomming")
    if not incident_log.empty:
        incident_log_dt = incident_log.copy()
        incident_log_dt['Date'] = pd.to_datetime(incident_log_dt['Date'])
        monthly_summary = incident_log_dt.groupby([pd.Grouper(key='Date', freq='M'), 'Category']).size().unstack(fill_value=0)
        monthly_summary.index = monthly_summary.index.strftime('%Y-%m')
        st.dataframe(monthly_summary.head(10), use_container_width=True, height=300)
        fig, ax = plt.subplots(figsize=(6, 3))
        monthly_summary.plot(
            kind='bar', 
            stacked=True, 
            ax=ax, 
            color=sns.color_palette('tab10', n_colors=len(monthly_summary.columns))
        )
        ax.set_title('Maandelikse Insidente', fontsize=10)
        ax.set_xlabel('Maand', fontsize=8)
        ax.set_ylabel('Aantal', fontsize=8)
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        ax.tick_params(axis='both', labelsize=7)
        ax.legend(title='Kategorie', fontsize=7)
        plt.tight_layout()
        st.pyplot(fig)
        plt.close()
    else:
        st.write("Geen insidente om te wys nie.")

with tab4:
    st.subheader("Kwartaallikse Opsomming")
    if not incident_log.empty:
        incident_log_dt = incident_log.copy()
        incident_log_dt['Date'] = pd.to_datetime(incident_log_dt['Date'])
        quarterly_summary = incident_log_dt.groupby([pd.Grouper(key='Date', freq='Q'), 'Category']).size().unstack(fill_value=0)
        quarterly_summary.index = quarterly_summary.index.map(
            lambda x: f"{x.year}-Q{(x.month-1)//3 + 1}"
        )
        st.dataframe(quarterly_summary.head(10), use_container_width=True, height=300)
        fig, ax = plt.subplots(figsize=(6, 3))
        quarterly_summary.plot(
            kind='bar', 
            stacked=True, 
            ax=ax, 
            color=sns.color_palette('tab10', n_colors=len(quarterly_summary.columns))
        )
        ax.set_title('Kwartaallikse Insidente', fontsize=10)
        ax.set_xlabel('Kwartaal', fontsize=8)
        ax.set_ylabel('Aantal', fontsize=8)
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        ax.tick_params(axis='both', labelsize=7)
        ax.legend(title='Kategorie', fontsize=7)
        plt.tight_layout()
        st.pyplot(fig)
        plt.close()
    else:
        st.write("Geen insidente om te wys nie.")

# High-risk learners
st.subheader("Leerders met Herhalende Insidente")
incident_counts = incident_log['Learner_Full_Name'].value_counts()
high_risk_learners = incident_counts[incident_counts > 2].index
high_risk_df = incident_log[incident_log['Learner_Full_Name'].isin(high_risk_learners)]

if not high_risk_df.empty:
    st.markdown("Leerders met meer as twee insidente:")
    display_df = high_risk_df.rename(columns={
        'Learner_Full_Name': 'Leerder Naam',
        'Class': 'Klas',
        'Teacher': 'Onderwyser',
        'Incident': 'Insident',
        'Category': 'Kategorie',
        'Comment': 'Kommentaar',
        'Date': 'Datum',
        'Sanction_Resolved': 'Sanksie Opgelos'
    })
    st.dataframe(
        display_df,
        use_container_width=True,
        column_config={
            "Leerder Naam": st.column_config.TextColumn("Leerder Naam", width="medium"),
            "Klas": st.column_config.TextColumn("Klas", width="small"),
            "Onderwyser": st.column_config.TextColumn("Onderwyser", width="medium"),
            "Insident": st.column_config.TextColumn("Insident", width="medium"),
            "Kategorie": st.column_config.TextColumn("Kategorie", width="small"),
            "Kommentaar": st.column_config.TextColumn("Kommentaar", width="large"),
            "Datum": st.column_config.DateColumn("Datum", width="medium", format="YYYY-MM-DD"),
            "Sanksie Opgelos": st.column_config.CheckboxColumn("Sanksie Opgelos", width="small")
        }
    )
else:
    st.info("Geen leerders met herhalende insidente nie.")
